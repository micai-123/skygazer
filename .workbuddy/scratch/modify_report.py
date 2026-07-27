# -*- coding: utf-8 -*-
import docx
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

SRC = r'D:/桌面/24软工一班-24050539110-孙瑞铭-2026.7.10/实训总结报告/（24050539110-孙瑞铭）skygazer天气预测与查询系统.docx'

doc = docx.Document(SRC)
body = doc.element.body

def text_of(el):
    return ''.join(t.text or '' for t in el.iter(qn('w:t')))

def make_para(text, align='left', size=12, bold=False, cn='宋体', en=None,
              before=0, after=0, line=18, page_break=False):
    """align: left/center/right ; size in pt ; before/after/line in pt"""
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), align); pPr.append(jc)
    sp = OxmlElement('w:spacing')
    sp.set(qn('w:before'), str(int(before*20)))
    sp.set(qn('w:after'), str(int(after*20)))
    sp.set(qn('w:line'), str(int(line*20)))
    sp.set(qn('w:lineRule'), 'exact')
    pPr.append(sp)
    p.append(pPr)
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts')
    rf.set(qn('w:eastAsia'), cn)
    rf.set(qn('w:ascii'), en or cn)
    rf.set(qn('w:hAnsi'), en or cn)
    rPr.append(rf)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size*2))); rPr.append(sz)
    szc = OxmlElement('w:szCs'); szc.set(qn('w:val'), str(int(size*2))); rPr.append(szc)
    if bold:
        rPr.append(OxmlElement('w:b'))
        rPr.append(OxmlElement('w:bCs'))
    r.append(rPr)
    t = OxmlElement('w:t'); t.set(qn('xml:space'), 'preserve'); t.text = text
    r.append(t)
    if page_break:
        br = OxmlElement('w:br'); br.set(qn('w:type'), 'page'); r.append(br)
    p.append(r)
    return p

def add_before(ref, p):
    body.insert(body.index(ref), p)

# ---- 1. 找到首个非空 body 子元素作为参考点 ----
ref = None
for ch in body.iterchildren():
    tag = ch.tag.split('}')[-1]
    if tag == 'p':
        if text_of(ch).strip():
            ref = ch; break
    elif tag == 'tbl':
        ref = ch; break
    # sectPr 跳过
assert ref is not None, '未找到参考元素'

# ---- 2. 删除开头的空段落（原封面残留脚手架） ----
to_remove = []
for ch in body.iterchildren():
    if ch is ref:
        break
    if ch.tag.split('}')[-1] == 'p' and not text_of(ch).strip():
        to_remove.append(ch)
    else:
        break
for el in to_remove:
    body.remove(el)
print('已删除开头空段落:', len(to_remove))

# ---- 3. 构建彭于晏组风格封面（组长孙瑞铭 / 组员李宇军）----
cover = []
cover.append(make_para('', after=12))                                   # 顶部留白
cover.append(make_para('安博教育集团宜春实训基地', align='left', size=16,
                       cn='宋体', after=10))                            # 顶部机构名(左)
cover.append(make_para('', after=18))                                  # 留白
cover.append(make_para('实训总结报告', align='center', size=24,
                       bold=True, cn='黑体', after=12))                 # 大标题
cover.append(make_para('', after=18))                                  # 留白
cover.append(make_para('实训 组长：孙瑞铭', align='center', size=15,
                       bold=True, cn='黑体', after=6))                  # 组长
cover.append(make_para('组员：李宇军', align='center', size=15,
                       bold=True, cn='黑体', after=10))                 # 组员
cover.append(make_para('', after=24))                                  # 留白
cover.append(make_para('安博教育集团宜春实训基地', align='center', size=15,
                       bold=True, cn='宋体', after=4))                  # 底部机构名
cover.append(make_para('2026年7月', align='center', size=15,
                       bold=True, cn='宋体', after=0, page_break=True)) # 日期 + 分页

# 按正确顺序插入到 ref 之前
for p in cover:
    add_before(ref, p)
print('封面已插入（%d 行）' % len(cover))

# ---- 4. 在 17.1 小组分工内、附录之前 补入李宇军的任务 ----
bullet_style = doc.styles['List Bullet']
appendix_ref = None
for p in doc.paragraphs:
    if p.text.strip() == '附录' and p.style.name == '一级标题':
        appendix_ref = p._p
        break
assert appendix_ref is not None, '未找到“附录”标题'

def add_para_before(ref_el, text, style=None, bold=False, size=12, cn='宋体'):
    p = doc.add_paragraph()
    if style is not None:
        p.style = style
    else:
        p.paragraph_format.line_spacing = None
    r = p.add_run(text)
    r.font.bold = bold
    r.font.size = docx.shared.Pt(size)
    r.font.name = cn
    rPr = r._element.get_or_add_rPr()
    rf = rPr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts'); rPr.append(rf)
    rf.set(qn('w:eastAsia'), cn)
    if rPr.find(qn('w:sz')) is None:
        sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(int(size*2))); rPr.append(sz)
    body.insert(body.index(ref_el), p._p)
    return p

li = []
li.append(add_para_before(appendix_ref, '李宇军：', bold=False, size=12, cn='宋体'))
li.append(add_para_before(appendix_ref,
    '本文所做的主要工作集中在后端天气数据查询、测试保障与前端呈现一致性三个方面，具体如下：',
    size=12, cn='宋体'))
li.append(add_para_before(appendix_ref,
    '（1）后端天气查询工作：基于 Spring Boot 与 MyBatis 实现天气数据查询接口，对接和风天气 API，完成实时天气、逐小时与七日预报、空气质量等数据的获取、清洗与持久化，并通过 Redis 缓存优化查询性能，支撑天气服务模块的核心数据链路；',
    style=bullet_style))
li.append(add_para_before(appendix_ref,
    '（2）测试单元编写工作：使用 JUnit 5 与 Mockito 编写后端单元测试，覆盖天气查询服务、缓存命中与异常降级等核心逻辑，保障关键业务在持续迭代中的稳定性与可维护性；',
    style=bullet_style))
li.append(add_para_before(appendix_ref,
    '（3）前端风格控制工作：统一前端 UI 设计规范与组件样式（Vue 3 + Element Plus / CSS 变量），控制天气卡片、图表与地图等页面的视觉风格与排版一致性，提升系统界面的美观度与用户体验。',
    style=bullet_style))
li.append(add_para_before(appendix_ref,
    '通过参与 SkyGazer 项目的后端天气查询、单元测试与前端风格控制工作，我进一步加深了对前后端协作流程与工程质量保障的理解，也在版本管理与团队配合中提升了协作与沟通能力。',
    size=12, cn='宋体'))
print('李宇军任务已补入（%d 段）' % len(li))

doc.save(SRC)
print('已保存:', SRC)
