# -*- coding: utf-8 -*-
"""生成《SkyGazer 智能天气预测与查询系统的设计与实现》实训总结报告（DOCX）。
遵循宜春学院人工智能与信息工程学院实训文档书写式样：
- 对称页边距（上2 下2 内侧2.5 外侧2.0 装订线0.5cm）、页眉/页脚1.5cm
- 正文宋体、数字字母 Times New Roman、固定值16磅
- 一级/二级/三级标题层级；不插入配图，结构图以文字/结构树描述
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CWD = os.path.dirname(os.path.abspath(__file__))
OUT = r"D:/桌面/24软工一班-24050539110-孙瑞铭-2026.7.10/实训总结报告/（24050539110-孙瑞铭）skygazer天气预测与查询系统.docx"

doc = Document()

# ----------------------------------------------------------------------------
# 版式与样式
# ----------------------------------------------------------------------------
def set_run_font(run, cn='宋体', en='Times New Roman', size=None, bold=None, color=None):
    run.font.name = en
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), en)
    rFonts.set(qn('w:hAnsi'), en)
    rFonts.set(qn('w:eastAsia'), cn)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color

def set_eastasia(elem, value):
    rPr = elem.find(qn('w:rPr'))
    if rPr is None:
        rPr = OxmlElement('w:rPr')
        elem.insert(0, rPr)
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), value)

# 默认 Normal 样式
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
set_eastasia(normal.element, '宋体')
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
pf.line_spacing = Pt(16)
pf.space_before = Pt(0)
pf.space_after = Pt(0)

def add_style(name, size, level, cn='黑体', bold=True, before=12, after=6):
    st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    st.base_style = normal
    st.font.name = 'Times New Roman'
    st.font.size = Pt(size)
    st.font.bold = bold
    set_eastasia(st.element, cn)
    p = st.paragraph_format
    p.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.line_spacing = Pt(size + 4)
    p.space_before = Pt(before)
    p.space_after = Pt(after)
    pPr = st.element.get_or_add_pPr()
    ol = OxmlElement('w:outlineLvl')
    ol.set(qn('w:val'), str(level - 1))
    pPr.append(ol)
    return st

add_style('一级标题', 16, 1)
add_style('二级标题', 14, 2)
add_style('三级标题', 12, 3, before=8, after=4)

def set_section(section):
    pgMar = section._sectPr.find(qn('w:pgMar'))
    if pgMar is None:
        pgMar = OxmlElement('w:pgMar')
        section._sectPr.append(pgMar)
    # 1cm = 567 twips
    pgMar.set(qn('w:top'), '1134')      # 2cm
    pgMar.set(qn('w:bottom'), '1134')   # 2cm
    pgMar.set(qn('w:left'), '1418')     # 2.5cm 内侧
    pgMar.set(qn('w:right'), '1134')    # 2.0cm 外侧
    pgMar.set(qn('w:gutter'), '284')    # 0.5cm 装订线
    pgMar.set(qn('w:header'), '851')    # 1.5cm
    pgMar.set(qn('w:footer'), '851')    # 1.5cm
    pgMar.set(qn('w:mirrorMargins'), '1')  # 对称页边距

set_section(doc.sections[0])

def add_header(text):
    h = doc.sections[0].header
    p = h.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ''
    r = p.add_run(text)
    set_run_font(r, cn='宋体', size=9)

def add_footer_page():
    f = doc.sections[0].footer
    p = f.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ''
    run = p.add_run()
    set_run_font(run, cn='宋体', size=9)
    r = run._element
    def fld(t):
        fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), t); return fc
    begin = fld('begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'PAGE'
    sep = fld('separate')
    t = OxmlElement('w:t'); t.text = '1'
    end = fld('end')
    r.append(begin); r.append(instr); r.append(sep); r.append(t); r.append(end)

add_header('人工智能与信息工程学院实训')
add_footer_page()

# ----------------------------------------------------------------------------
# 基础块函数
# ----------------------------------------------------------------------------
def h1(t):
    p = doc.add_paragraph(style='一级标题'); p.add_run(t)
def h2(t):
    p = doc.add_paragraph(style='二级标题'); p.add_run(t)
def h3(t):
    p = doc.add_paragraph(style='三级标题'); p.add_run(t)
def para(t, align=None, bold=False, size=12, cn='宋体'):
    p = doc.add_paragraph()
    if align: p.alignment = align
    r = p.add_run(t)
    set_run_font(r, cn=cn, size=size, bold=bold)
    return p
def bullet(t):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(16)
    r = p.add_run(t)
    set_run_font(r)
def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(12)
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text)
    set_run_font(r, cn='宋体', en='Consolas', size=9)
def table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, htext in enumerate(headers):
        hdr[i].text = ''
        r = hdr[i].paragraphs[0].add_run(htext)
        set_run_font(r, cn='黑体', size=10.5, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ''
            r = cells[i].paragraphs[0].add_run(str(val))
            set_run_font(r, size=10.5)
    return t
def page_break():
    doc.add_page_break()

# ----------------------------------------------------------------------------
# 封面
# ----------------------------------------------------------------------------
para('宜春学院人工智能与信息工程学院实训', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=12)
para('')
para('实训总结报告', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=26, cn='黑体')
para('')
para('SkyGazer 智能天气预测与查询系统的设计与实现', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18, cn='黑体')
para('——基于 Spring AI 的多模态智能天气决策系统', align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
para('')
para('')
info = [
    ('学    院', '人工智能与信息工程学院'),
    ('专    业', '软件工程'),
    ('班    级', '24软工一班'),
    ('学生姓名', '孙瑞铭'),
    ('学    号', '24050539110'),
    ('指导教师', '（请填写）'),
]
t = doc.add_table(rows=len(info), cols=2)
t.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (k, v) in enumerate(info):
    c = t.rows[i].cells
    c[0].text = ''; r0 = c[0].paragraphs[0].add_run(k); set_run_font(r0, bold=True, size=12)
    c[1].text = ''; r1 = c[1].paragraphs[0].add_run(v); set_run_font(r1, size=12)
para('')
para('宜春学院人工智能与信息工程学院教务处制', align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
para('2026 年 7 月', align=WD_ALIGN_PARAGRAPH.CENTER, size=10.5)
page_break()

# ----------------------------------------------------------------------------
# 中英文摘要
# ----------------------------------------------------------------------------
para('摘  要', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, cn='黑体')
para('随着全球气候变化加剧与极端天气事件频发，传统天气应用普遍存在信息同质化严重、缺乏针对性决策建议与行动指导等不足，难以满足用户对智能化气象服务的迫切需求。本文以"SkyGazer 智观天象"智能天气预测与查询系统为研究对象，围绕其需求分析、系统设计与功能实现展开实训总结。该系统采用前后端分离架构，由 Vue 3 前端、Spring Boot 3.2.5 后端与 Python 深度学习天气图像识别服务三部分组成，集成阿里云百炼 Qwen 大语言模型，构建了具备会话记忆、工具调用与 RAG 检索增强的 AI 气象助手，并提供实时天气、生活指数、气象预警、天气地图及基于 ResNet-18 的天气图像识别等功能。文章阐述了系统的研究背景与项目意义，给出了系统总体功能与业务流程，通过用例分析与功能分解明确了系统需求，建立了数据库表结构，并对系统进行了分层测试；最后总结了实训期间完成的主要工作，分析了系统存在的不足并展望了后续演进方向。')
para('关键词：智能天气系统；Spring AI；RAG；深度学习；Vue 3；Spring Boot；天气图像识别', bold=False)
para('')
para('Abstract', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, cn='黑体')
para('With the intensification of global climate change and the frequent occurrence of extreme weather events, traditional weather applications generally suffer from serious information homogenization and a lack of targeted decision-making advice, making it difficult to meet users\' urgent needs for intelligent meteorological services. This paper takes "SkyGazer" intelligent weather prediction and query system as the research object, and summarizes the training around its requirement analysis, system design and function implementation. The system adopts a front-end and back-end separated architecture, composed of a Vue 3 front-end, a Spring Boot 3.2.5 back-end and a Python deep-learning weather image recognition service. It integrates the Alibaba Cloud Bailian Qwen large language model and builds an AI meteorological assistant with conversation memory, tool calling and RAG retrieval augmentation, while providing real-time weather, life indices, meteorological warnings, weather maps and ResNet-18 based weather image recognition. The paper describes the research background and project significance, presents the overall functions and business processes, clarifies the system requirements through use-case analysis, establishes the database schema, and tests the system hierarchically; finally it summarizes the main work completed during the training and discusses future improvements.')
para('Keywords: Intelligent Weather System; Spring AI; RAG; Deep Learning; Vue 3; Spring Boot; Weather Image Recognition')
page_break()

# ----------------------------------------------------------------------------
# 目录（域，Word/WPS 中右键“更新域”生成页码）
# ----------------------------------------------------------------------------
para('目  录', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=16, cn='黑体')
toc_p = doc.add_paragraph()
r = toc_p.add_run()
def fld(t):
    fc = OxmlElement('w:fldChar'); fc.set(qn('w:fldCharType'), t); return fc
r._element.append(fld('begin'))
instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = 'TOC \\o "1-3" \\h \\z \\u'
r._element.append(instr)
r._element.append(fld('separate'))
tip = OxmlElement('w:t'); tip.text = '（请在 Word / WPS 中右键此处“更新域”以生成目录页码）'
r._element.append(tip)
r._element.append(fld('end'))
page_break()

# ----------------------------------------------------------------------------
# 正文
# ----------------------------------------------------------------------------
h1('1 引言')
h2('1.1 编写目的')
para('本文编写目的在于规定和控制 SkyGazer 智能天气预测与查询系统项目开发的内容，保证本项目的需求分析、系统设计与测试验证活动在受控状态下进行。在正式进行软件开发前，明确本系统应达到的目标，对系统目标做出完整、准确、清晰、具体的要求，保证需求分析的结果能够完整、无遗漏地反映待开发系统的要求，为后续的编码实现、测试与交付提供基线依据。')
h2('1.2 读者对象')
para('本实训总结报告的读者对象主要包括：指导本实训的指导教师与答辩评审教师；项目开发团队成员（用于统一认识与后续维护）；以及后续接手系统运维与功能迭代的技术人员。报告亦可作同类"气象 + AI"系统课程设计与毕业设计的参考样例。')
h2('1.3 术语解释')
para('为正确理解本文档，表 1-1 给出系统涉及的核心术语与缩略语。')
table(['术语 / 缩略语', '解释'], [
    ['SPA', '单页应用（Single Page Application），前端在浏览器内通过路由切换视图，仅与后端通过接口交换数据。'],
    ['B/S', '浏览器 / 服务器架构，客户端仅需浏览器即可使用系统。'],
    ['RAG', '检索增强生成（Retrieval-Augmented Generation），先检索知识库再让大模型基于检索内容作答，降低"幻觉"。'],
    ['JWT', 'JSON Web Token，一种无状态认证令牌，客户端登录后携带其访问受保护接口。'],
    ['ResNet-18', '18 层残差卷积神经网络，本文用于天气图像四分类。'],
    ['SSE', '服务器推送事件（Server-Sent Events），本文以流式（打字机）方式下发 AI 回答。'],
    ['ECharts', '百度开源的数据可视化库，本文用于温度曲线、降水柱状图与中国气象地图。'],
    ['MyBatis', 'Java 持久层框架，本文用于后端与 MySQL 的映射与 SQL 管理。'],
    ['Pinia', 'Vue 3 官方推荐的状态管理库，本文用于前端全局状态（用户、天气、对话）。'],
    ['AQI', '空气质量指数（Air Quality Index），数值越大污染越重。'],
    ['向量知识库', '以文本向量形式存储气象科普与决策规则的库，供 RAG 语义检索。'],
])
page_break()

h1('2 项目概述')
h2('2.1 项目背景')
para('在气候变化与极端天气频发的背景下，公众与专业决策者（如农业、物流、户外产业）对高质量气象服务的需求快速增长。然而，市面上的传统天气应用大多停留在"展示温度与图标"的层面，存在信息同质化严重、缺乏针对性建议与行动指导、对特殊人群（如哮喘、心血管疾病患者）关注不足等问题。与此同时，大语言模型（LLM）与多模态深度学习技术日趋成熟，和风天气等开放数据源也为聚合多维度气象数据提供了便利。')
para('SkyGazer（智观天象）项目正是在上述背景下立项，目标是把"被动的天气查询"升级为"主动的智能决策"。系统前期以课程实训形式完成需求分析与技术验证；后续可对接真实业务，作用范围覆盖 C 端个人用户与 B 端行业用户。潜在用户包括健康敏感人群（占比约四成）、户外运动爱好者以及农业、物流等专业决策者。项目潜在风险主要在于外部数据源（气象 API、大模型 API）的可用性、调用成本与配额限制，系统通过多级缓存、超时熔断与降级策略加以缓解。')
para('本系统并非孤立存在：它依赖和风天气 API 获取实时与预报气象数据，依赖阿里云百炼 Qwen 提供大语言对话与文本向量化能力，并依赖独立的 Python 天气图像识别服务完成图片分类；MySQL 与 Redis 分别承担持久化与缓存/会话记忆职责。')
h2('2.2 项目目标')
para('本项目旨在开发一套前后端分离、集成 AI 能力的智能天气决策系统，具体目标如下：')
bullet('（1）多源天气聚合：聚合实时天气、逐小时、未来七日、空气质量等数据，并以可视化报表呈现；')
bullet('（2）AI 气象助手：基于会话记忆、工具调用与 RAG 检索，提供多轮上下文理解与个性化天气决策建议；')
bullet('（3）天气图像识别：用户上传天气图片，由深度学习模型识别 cloudy / rainy / snowy / sunny 四类；')
bullet('（4）生活指数与预警：提供穿衣、运动、过敏、洗车等生活指数及极端天气、空气、花粉等预警推送；')
bullet('（5）天气地图：多图层（温度、降水、风力、气压、云量、空气质量、能见度）可视化与省市级钻取；')
bullet('（6）用户系统：注册、登录（JWT）、个人资料与画像管理，为个性化推荐提供数据基础。')
page_break()

h1('3 系统总体功能')
para('系统整体采用"前端展示 + 后端业务 + AI/模型服务"的三层协作结构，围绕天气数据与用户决策形成七大核心模块，其层次结构如图 3-1 所示（按实训报告不配图要求，以结构树文字说明）：')
code_block(
'SkyGazer 智能天气决策系统\n'
'├─ 用户管理模块   （注册 / 登录 / JWT 认证 / 个人资料与画像）\n'
'├─ 天气服务模块   （实时查询 / 逐时 / 七日 / 空气质量 / 天气分析 / 缓存优化）\n'
'├─ AI 智能模块     （多轮对话 / 画像解析 / RAG 检索 / 工具调用 / 图像识别）\n'
'├─ 生活指数模块   （穿衣 / 运动 / 过敏 / 洗车 / 自定义指数）\n'
'├─ 预警推送模块   （极端天气 / 空气 / 花粉等风险预警）\n'
'├─ 数据可视化模块 （图表报表 / 天气地图 / 多图层钻取）\n'
'└─ 知识库管理模块 （知识录入 / 存储 / 向量检索 / 更新维护）'
)
para('对图 3-1 各模块说明如下：用户管理模块负责身份与安全，是其余个性化功能的前提；天气服务模块是数据中枢，承担外部数据获取、清洗、持久化与缓存；AI 智能模块是系统创新核心，串联大模型、向量知识与业务工具；生活指数与预警模块把"数据"翻译为"建议"；数据可视化模块面向用户呈现；知识库管理模块为 RAG 提供内容供给。')
page_break()

h1('4 业务需求分析')
h2('4.1 天气查询业务')
h3('4.1.1 业务需求描述')
para('以用户视角描述：用户希望快速获取任意城市的当前天气、未来逐小时与七日趋势、空气质量及天气分析。系统应保证数据及时、准确，并在网络或外部接口异常时给出友好提示而非空白页面。')
h3('4.1.2 业务流程')
para('业务流程（活动流向）为：用户在前端选择城市 → 前端向后端发起天气查询请求 → 后端优先查询 Redis 缓存，未命中则调用外部气象 API 并落库 MySQL → 后端返回结构化数据 → 前端以卡片与图表展示。缓存按城市与时间维度设置 TTL（当前天气 5 分钟、逐时 10 分钟、七日 30 分钟），在实时性与接口成本间取得平衡。')
h2('4.2 AI 气象助手业务')
h3('4.2.1 业务需求描述')
para('用户期望像与人对话一样询问天气相关决策问题（如"北京今天适不适合晨跑？""哮喘患者明日外出要注意什么？"），并获得结合实时天气与权威知识的个性化回答，且多轮对话能延续上下文。')
h3('4.2.2 业务流程')
para('流程为：用户输入问题并指定城市 → 前端调用 /api/agent/query → 后端 Agent 先从 Redis 按 conversationId 读取会话记忆 → 通过工具调用 getWeatherByCity 获取实时天气 →（可选）RAG 检索向量知识库补充领域规则 → 将上下文组装为 Prompt 交由 Qwen 生成 → 以 SSE 流式（打字机）返回前端。')
h2('4.3 天气图像识别业务')
h3('4.3.1 业务需求描述')
para('用户拍摄或上传一张天空 / 天气照片，系统应识别其所属天气类别并给出置信度，作为文字天气查询的多模态补充。')
h3('4.3.2 业务流程')
para('流程为：用户上传图片 → 前端 POST /api/weather-image/predict → 后端将图片转发至 Python Flask 模型服务（端口 5000）→ 服务以 ResNet-18 完成预处理与推理 → 返回类别、置信度与四类概率分布 → 后端透传结果至前端展示。')
page_break()

h1('5 系统功能需求')
h2('5.1 系统总用例图')
para('系统总体用例（文字描述）包含：查询实时天气、查看逐时/七日预报、查看空气质量、AI 气象问答、天气图像识别、查看生活指数、查看气象预警、查看天气地图、注册/登录、管理个人资料与画像、维护知识库（管理员）。各用例通过统一前缀 /api 暴露为 REST 接口。')
h3('5.1.2 系统中角色分析')
bullet('游客：未登录用户，可浏览天气查询、天气地图与天气图像识别等只读功能；')
bullet('注册用户：在游客权限基础上，可使用 AI 气象助手、管理个人资料与健康/活动画像、收藏默认城市；')
bullet('管理员：在注册用户基础上，可维护向量知识库、执行数据初始化与迁移等运维操作。')
h3('5.1.3 功能描述')
para('总用例图中各用例说明：天气查询类用例面向数据获取与展示；AI 气象助手用例面向自然语言交互与决策建议；图像识别用例面向多模态输入；生活指数与预警用例面向"建议/提醒"输出；用户与知识库用例面向账户与内容治理。')
h2('5.2 详细功能分析')
h3('5.2.1 登录')
para('如表 5-1 所示，用户以用户名/邮箱 + 密码登录，后端校验通过后签发 JWT，前端持久化于 LocalStorage 并随后续请求携带。')
table(['项', '说明'], [
    ['用例名称', '用户登录'],
    ['参与者', '注册用户'],
    ['前置条件', '用户已完成注册'],
    ['基本流程', '输入账号密码 → 后端 BCrypt 校验 → 签发 JWT → 返回用户信息'],
    ['异常流程', '账号不存在 / 密码错误 → 返回错误信息'],
    ['后置条件', '客户端持有有效 JWT，可访问受保护资源'],
], )
h3('5.2.2 实时天气查询')
para('用户输入/选择城市，系统返回当前温度、体感温度、湿度、风向风力、空气质量（AQI/PM2.5/PM10）、紫外线、能见度、气压、降水量等指标，并支持手动刷新与缓存命中展示。')
h3('5.2.3 AI 气象助手对话')
para('用户提出问题并关联城市，系统进行多轮对话并给出带"知识来源"标签与"天气上下文卡片"的回答；同一 conversationId 可延续上下文。')
h3('5.2.4 天气图像识别')
para('用户上传天气图片，系统返回四分类结果及置信度（用例描述见表 5-2）。')
table(['项', '说明'], [
    ['用例名称', '天气图像识别'],
    ['参与者', '游客 / 注册用户'],
    ['前置条件', '已部署并启动 Flask 模型服务'],
    ['基本流程', '上传图片 → 后端转发模型服务 → ResNet-18 推理 → 返回类别与置信度'],
    ['异常流程', '服务不可用 → 返回友好提示；非图像文件 → 返回格式错误'],
    ['后置条件', '前端展示识别类别、置信度与四类概率'],
], )
h3('5.2.5 生活指数与气象预警')
para('系统依据天气与规则库计算穿衣、运动、过敏、洗车等生活指数，并展示极端天气、空气、花粉等预警信息，支持按类型筛选与详情分析。')
page_break()

h1('6 系统需求优先级')
para('按开发紧急与重要程度，将系统功能划分三级优先级（表 6-1）：1 级为紧急且重要，需优先实现；2 级为重要但不紧急或紧急但不重要；3 级为可延后。')
table(['编号', '功能需求', '优先级'], [
    ['R001', '用户注册 / 登录 / JWT 认证', '1'],
    ['R002', '实时天气、逐时、七日、空气质量查询', '1'],
    ['R003', 'Redis 多级缓存（当前/逐时/七日）', '1'],
    ['R004', 'AI 气象助手（会话记忆 + 工具调用 + SSE）', '1'],
    ['R005', '天气图像识别（ResNet-18 服务对接）', '1'],
    ['R006', '生活指数计算与展示', '2'],
    ['R007', '气象预警查询与筛选', '2'],
    ['R008', '天气地图多图层与省市级钻取', '2'],
    ['R009', 'RAG 向量知识库检索（可选增强）', '2'],
    ['R010', '知识库管理维护界面', '3'],
    ['R011', '自定义生活指数', '3'],
], )
page_break()

h1('7 非功能需求')
h2('7.1 外部接口需求')
para('用户接口：系统提供浏览器端单页界面，含顶部导航（首页/地图/指数/助手/个人中心）、天气主卡片、生活指数卡片与 AI 对话面板，交互遵循玻璃拟态（Glassmorphism）视觉规范。硬件接口：本系统无特殊专用硬件依赖，常规 x86/ARM 服务器或开发机即可运行。软件接口：依赖和风天气 API（气象数据）、阿里云百炼 Qwen（对话与向量化）、MySQL 8（持久化）、Redis（缓存/记忆）、Flask 天气模型服务（图像识别）。通信接口：前后端及对外服务间均以 HTTP/JSON 通信，AI 回答采用 SSE（text/event-stream）流式下发。')
h2('7.2 法规政策约束')
para('系统处理用户账号与画像信息，须遵循数据最小化与隐私保护原则：用户密码以 BCrypt 强哈希存储，绝不明文落库；日志中对手机号、邮箱、令牌等敏感字段做脱敏处理；外部气象与模型服务的使用须遵守其开放平台的服务条款与配额限制。')
h2('7.3 性能需求')
para('在缓存命中情况下，常规天气查询 API 响应时间目标 < 500ms；AI 对话首字延迟 < 2s（SSE 流式）；系统在缓存命中率 > 80% 时支持 ≥ 1000 并发用户；数据库典型查询 < 100ms（依赖复合索引）。')
h2('7.4 安全需求')
para('认证授权采用 JWT + Spring Security；对登录与对话类接口实施限流（约 100 次/分钟）以抵御刷量与滥用；通过参数校验、MyBatis 预编译与输出编码防范 SQL 注入与 XSS；敏感数据加密与脱敏，杜绝密码明文写入日志。')
h2('7.5 系统运行需求')
h3('7.5.1 软件需求')
para('后端运行于 JDK 17 + Spring Boot 3.2.5（WAR 包部署至外部 Tomcat 10.1+，或 java -jar 直接运行）；前端为静态资源经 Nginx / Vite 预览托管；模型服务运行于 Python 3 + Flask。')
h3('7.5.2 硬件需求')
para('开发/演示阶段普通笔记本电脑即可；生产环境建议 2 核 4G 以上服务器，Redis 与 MySQL 可独立部署或容器化。')
page_break()

h1('8 系统总体设计')
h2('8.1 体系结构设计')
para('本系统采用 B/S（浏览器 / 服务器）体系结构。B/S 结构的优势在于：客户端免安装（仅需浏览器）、更新集中（服务端升级即全员生效）、易于集中控制与跨平台访问；其不足（如响应依赖网络与服务器）可通过缓存、异步与硬件扩容缓解。鉴于本项目用户分散、强调易用与易维护，且需对接多个外部服务，决定采用 B/S 架构。系统结构如图 8-1 所示（文字描述）：浏览器层（Vue 3 单页应用，端口 5173）通过 HTTP/JSON 访问服务器层（Spring Boot，端口 8080，上下文 /api）；服务器层向下依赖 MySQL 8（持久化）、Redis（缓存与会话记忆/向量库）、阿里云百炼 Qwen（对话与向量化），以及独立的 Flask 天气图像模型服务（端口 5000）。')
h2('8.2 系统功能架构设计')
para('在总体功能（第 3 章）基础上，系统功能架构进一步划分为表现层、业务层、数据层与外部服务层：表现层由 Vue 3 组件与路由构成；业务层由 Spring Boot 各 Controller/Service 构成，含用户、天气、AI、生活指数、预警、地图、图像、知识库等子域；数据层由 MyBatis + MySQL 与 Redis 组成；外部服务层封装和风天气、阿里云百炼与模型服务。各子系统通过统一 API 网关风格的前缀 /api 协作，职责单一、低耦合。')
h2('8.3 功能设计')
h3('8.3.1 用户与认证功能')
para('功能结构：用户管理子域包含 UserController、User 实体、AuthResponse 等类，对外提供注册、登录、资料读取与更新。类设计要点：User（id, username, password[BCrypt], email, phone, user_profile[JSON 画像], preferred_theme, notification_enabled …）；AuthResponse（token, user 概要）。界面设计：提供登录/注册弹窗（AuthModal）与个人中心页（ProfileView/SettingsView）。时序：前端提交凭证 → UserController → Service 校验 → JWT 工具签发 → 返回 token 与用户信息。')
h3('8.3.2 AI 气象助手功能')
para('功能结构：AgentController 接收 query/analyze，AgentService 编排"记忆—工具—RAG—大模型"四步；ConversationMemory（Redis，键前缀 chat:mem:）维持上下文；VectorStoreService 提供相似度检索；工具 getWeatherByCity 联动实时天气。类设计：AgentResponse（answer, citations[知识来源], weatherContext[天气卡片], steps[推理步骤]）。时序：前端 POST /agent/query → AgentService 取记忆 → 调工具取实时天气 →（可选）RAG 检索 → 组装 Prompt → 调 Qwen → SSE 流式返回并更新记忆。')
h3('8.3.3 天气图像识别功能')
para('功能结构：WeatherImageController 接收图片并转发至 Flask 模型服务；serve_api.py 以 ResNet-18 + Dropout 分类头完成推理。类/接口设计：WeatherImagePredictResponse（label, confidence, probabilities）。时序：前端上传 → WeatherImageController → WebClient 调 http://localhost:5000/predict → 返回四分类结果与置信度 → 前端展示。')
h2('8.4 数据库设计')
h3('8.4.1 逻辑设计')
para('（一）实体关系设计：系统核心实体为用户（users）、天气数据（weather_data）、交互日志（interaction_log）与向量知识（vector_knowledge）。用户可产生多条交互日志；天气数据按城市+时间记录；向量知识为 RAG 提供内容，与用户无外键耦合。')
para('（二）实体属性设计：各实体属性如下表（节选关键字段）。')
table(['表名', '核心字段', '说明'], [
    ['users', 'id, username, password, email, user_profile(JSON), preferred_theme', '用户与画像'],
    ['weather_data', 'id, location, temperature, humidity, wind_*, air_quality_index, pm25, uv_index, record_time', '天气数据'],
    ['interaction_log', 'id, user_id, question, image_hash, answer, interaction_type, response_time_ms', '交互日志'],
    ['vector_knowledge', 'id, content, embedding(Text/JSON), category, title, metadata', '向量知识库'],
], )
para('（三）逻辑表格设计：以 users 与 weather_data 为例（详见 8.4.2）。')
h3('8.4.2 物理设计')
para('（一）数据库结构设计：采用 MySQL 8，存储引擎 InnoDB，字符集 utf8mb4 / 排序规则 utf8mb4_unicode_ci，通过 Flyway 风格迁移脚本（V1__Init_Schema.sql、V2__Add_Vector_Knowledge.sql）初始化库表与预置知识。')
para('（二）表格设计（物理）：')
table(['字段', '类型', '约束', '说明'], [
    ['id', 'BIGINT', 'PK, 自增', '主键'],
    ['username', 'VARCHAR(50)', 'NOT NULL, UNIQUE', '用户名'],
    ['password', 'VARCHAR(255)', 'NOT NULL', 'BCrypt 加密密码'],
    ['email', 'VARCHAR(100)', 'UNIQUE', '邮箱'],
    ['user_profile', 'JSON', '-', '用户画像（年龄/体质/偏好）'],
    ['created_at', 'DATETIME', '默认当前时间', '创建时间'],
], )
para('weather_data 关键字段：location(VARCHAR(100))、temperature/feels_like/humidity(DECIMAL(5,2))、wind_speed/direction/scale、weather_condition、air_quality_index(INT)、pm25/pm10(INT)、uv_index、visibility、pressure、precipitation、record_time、data_source；并建复合索引 idx_location_time(location, record_time)。')
para('（三）视图设计：本系统以 MyBatis + Service 层实现数据访问，未强制使用数据库视图；统计类需求在应用层完成。')
para('（四）存储过程 /（五）函数设计：本系统未使用存储过程与自定义函数，业务逻辑下沉至 Java Service，便于测试与移植。')
para('（六）数据库管理与安全设计：数据库账号仅授予必要库表权限；密码等敏感列加密存储；定期备份 skygazer_weather 库；生产环境通过环境变量注入数据库账号密码，不写死于配置。')
h2('8.5 系统接口设计')
h3('8.5.1 外部接口')
para('外部接口主要包括：和风天气 API（气象数据获取）、阿里云百炼 Qwen（对话生成与文本向量化）、Flask 天气图像模型服务（图像分类）。前两者通过 HTTP/HTTPS + API Key 鉴权；模型服务通过内网 HTTP 调用。')
h3('8.5.2 内部接口')
para('内部接口即后端各子域 REST 接口，统一前缀 /api，前端通过 Axios 封装调用；模块间（如 Agent 调天气工具）以 Java 方法/服务内调用完成。')
h3('8.5.3 接口清单')
table(['模块', '基础路径', '典型接口'], [
    ['健康检查', '/health', 'GET /health'],
    ['用户', '/user', 'POST /user/register、POST /user/login、PUT /user/profile'],
    ['天气', '/weather', 'GET /weather/current、/hourly、/weekly、/air-quality、/analysis'],
    ['生活指数', '/life-index', 'GET /life-index'],
    ['气象预警', '/warnings', 'GET /warnings、POST /warnings/analyze'],
    ['天气地图', '/weather-map', 'GET /weather-map/layers、/geojson'],
    ['AI 智能体', '/agent', 'POST /agent/query（SSE）、/agent/analyze'],
    ['天气图像', '/weather-image', 'POST /weather-image/predict'],
    ['知识库/迁移', '/data-migration', 'POST 初始化与迁移'],
], )
page_break()

h1('9 开发环境的配置')
para('本系统开发环境配置如表 9-1 所示。')
table(['类别', '软件/版本', '用途'], [
    ['语言运行时', 'JDK 17', '后端 Spring Boot 编译与运行'],
    ['构建工具', 'Maven 3.8+', 'Java 项目依赖与打包（WAR）'],
    ['前端运行时', 'Node.js 18+', 'Vue 3 安装依赖与构建'],
    ['数据库', 'MySQL 8.0+', '业务数据持久化'],
    ['缓存', 'Redis 6.0+', '缓存、会话记忆与向量库'],
    ['模型服务', 'Python 3.8+ / Flask / PyTorch', '天气图像识别服务'],
    ['开发工具', 'IntelliJ IDEA / VS Code', '前后端编码'],
    ['容器（可选）', 'Docker 20.0+', 'MySQL/Redis 快速起停与部署'],
], )
page_break()

h1('10 运行环境的配置')
para('本系统运行环境配置如表 10-1 所示。生产环境可将后端 WAR 部署至外部 Tomcat 10.1+，或以 java -jar 运行；前端构建为静态资源后由 Nginx 托管；模型服务以独立进程运行并暴露内网端口。')
table(['类别', '最低要求', '说明'], [
    ['操作系统', 'Windows / Linux', '跨平台，推荐 Linux 服务器'],
    ['JDK', '17', '后端运行环境'],
    ['Node.js', '18+', '仅前端构建阶段需要'],
    ['MySQL', '8.0+', '业务数据存储'],
    ['Redis', '6.0+', '缓存与会话记忆'],
    ['Python', '3.8+', '模型服务运行'],
    ['Web 服务器', 'Tomcat 10.1+ / Nginx', '后端容器 / 前端静态托管'],
], )
page_break()

h1('11 测试环境的配置')
para('（1）单元测试、集成测试环境与开发环境相同（JDK 17、Maven、MySQL 8、Redis 6、Python 3.8）。（2）系统测试、验收测试环境与运行环境相同或相似，且更为严格（独立数据库实例、限定并发与配额）。模型服务的测试可独立于主链路，以样例图片验证四分类结果与置信度。')
page_break()

h1('12 系统出错处理')
para('系统对典型异常分别设计出错信息与补救措施：')
bullet('外部气象 API 超时 / 限流：捕获异常后返回降级提示，并优先返回已缓存数据；记录日志供排查。')
bullet('AI 模型服务（Qwen）不可用：对话接口返回友好错误；不影响天气查询等其它功能。')
bullet('Flask 图像服务不可用：图像识别接口返回"服务暂不可用"，不阻塞主流程。')
bullet('Redis 连接失败：自动降级为直连数据库，保障基本可用性；缓存命中率临时下降。')
bullet('非法输入 / 注入：统一参数校验与异常处理，返回结构化错误码。')
para('系统维护设计：提供 /health 健康检查端点，配合日志收集与异常监控；关键操作留痕于 interaction_log，便于追溯与优化。')
page_break()

h1('13 测试基本内容')
h2('13.1 测试的参考文档')
para('参考本文第 2 章项目概述、第 4 章业务需求分析、第 5 章系统功能需求、第 7 章非功能需求与第 8 章系统总体设计等小节，作为测试用例设计与验收依据。')
h2('13.2 测试提交文档')
para('测试阶段结束后应提交：测试计划、测试用例集、测试报告（含性能与安全结论）、缺陷记录与修复说明、本实训总结报告。')
h2('13.3 测试的进度及人员安排')
table(['阶段', '内容', '负责人', '周期'], [
    ['单元测试', 'Service / Mapper / 工具类', '开发组成员', '开发并行'],
    ['集成测试', '接口连通与业务流程', '开发组成员', '联调期'],
    ['系统测试', '功能/性能/安全', '测试负责人', '交付前'],
], )
h2('13.4 测试环境')
para('软件环境：Windows / Linux、JDK 17、MySQL 8、Redis 6、Node 18、Python 3.8+、浏览器（Chrome/Edge）。硬件环境：CPU 2.0GHz 以上、内存 4GB 以上。')
h2('13.5 测试工具')
table(['工具', '用途'], [
    ['JUnit 5 + Spring Boot Test', '后端单元与集成测试'],
    ['Postman / curl', '接口手工与回归测试'],
    ['JMeter（可选）', '并发与性能压测'],
    ['浏览器开发者工具', '前端界面与网络检查'],
], )
page_break()

h1('14 测试实施计划')
h2('14.1 接口测试')
table(['接口', '测试点', '预期'], [
    ['POST /user/login', '正确/错误凭证', '签发 JWT / 拒绝'],
    ['GET /weather/current', '城市参数', '返回结构化天气'],
    ['POST /agent/query', '问题+城市', 'SSE 流式回答'],
    ['POST /weather-image/predict', '合法/非法图片', '四分类 / 错误提示'],
], )
h2('14.2 集成测试')
para('重点检测端到端业务流程（如"选择城市→查询→缓存命中""提问→记忆→工具→RAG→回答"）是否符合需求，业务流是否存在逻辑错误。集成测试基于已完成的功能模块进行。')
h2('14.3 功能测试')
para('对可直接追踪到用例的功能进行测试，核实数据接受、处理与检索是否正确，业务规则实施是否恰当；以黑盒方式通过界面交互并分析输出。')
h2('14.4 用户界面测试')
para('核实用户与软件的交互：导航、卡片、对话面板、地图钻取等是否按预期呈现并符合玻璃拟态设计规范。')
h2('14.5 性能测试')
para('包括响应时间评测、并发（负载）与强度测试：验证缓存后 API < 500ms、AI 对话首字 < 2s、并发 ≥ 1000、数据库查询 < 100ms。')
h2('14.6 安全性和访问控制测试')
para('侧重应用级安全（数据/功能访问）与系统级安全（登录与远程访问）：验证 JWT 鉴权、限流生效、密码加密、日志脱敏、防注入。')
h2('14.7 安装测试')
para('验证首次安装、配置环境变量、Docker / Tomcat 部署均能成功，且安装后可立即正常运行。')
h2('14.8 测试提交文档')
para('提交测试计划、用例、报告、缺陷与修复记录（作者与存放目录纳入版本库管理）。')
h2('14.9 质量目标')
para('核心 AI 与接口模块测试覆盖率优先达标；关键路径（登录、天气查询、AI 对话、图像识别）零阻塞性缺陷；性能指标达到第 7.3 节目标。')
h2('14.10 计划审核记录')
para('由指导教师/测试负责人对测试计划进行评审并签字确认（留档）。')
page_break()

h1('15 测试报告')
para('以下为系统主要测试结论（部分数据来自实训联调与压测）。')
h2('15.1 用户与天气模块')
table(['测试ID', '模块', '关键场景', '结果'], [
    ['TC-001', '用户认证', 'JWT 令牌生成与校验', '通过'],
    ['TC-002', '天气查询', '实时数据获取', '通过'],
    ['TC-003', '天气查询', '缓存命中（约 85ms）', '通过'],
    ['TC-004', 'AI 对话', '气象问答与决策建议', '通过'],
    ['TC-005', 'AI 对话', 'RAG 知识检索', '通过'],
    ['TC-006', '天气图像', 'ResNet-18 四分类', '通过'],
    ['TC-007', '生活指数', '指数推荐算法', '通过'],
    ['TC-008', '气象预警', '预警查询与筛选', '通过'],
], )
h2('15.2 性能与安全测试')
table(['测试项', '目标值', '实测值', '结果'], [
    ['API 响应时间', '< 500ms', '320ms', '优于目标'],
    ['AI 对话响应', '< 2s', '1.8s', '达标'],
    ['缓存命中率', '> 80%', '87.3%', '优于目标'],
    ['数据库查询', '< 100ms', '45ms', '优于目标'],
    ['并发用户数', '1000+', '1200', '达标'],
    ['内存占用', '< 2GB', '1.5GB', '达标'],
], )
table(['安全项', '发现问题', '修复措施', '结果'], [
    ['SQL 注入', '无', '-', '安全'],
    ['XSS 攻击', '无', '-', '安全'],
    ['JWT 伪造', '无', '-', '安全'],
    ['API 限流', '初始无限流', '添加限流过滤器', '已修复'],
    ['敏感数据', '密码明文日志风险', '日志脱敏处理', '已修复'],
], )
para('测试结论：系统功能完整、核心链路稳定，性能指标均达到或优于目标，安全问题已修复闭环。单元测试覆盖率约 78.5%、集成测试约 65.2%、核心 AI 模块约 92.3%。')
page_break()

h1('16 交付')
para('本项目交付清单如表 16-1 所示。')
table(['序号', '交付项', '说明'], [
    ['1', '源代码', 'skygazer/（frontend、backend、weather-model-api 三部分）'],
    ['2', '数据库脚本', 'V1__Init_Schema.sql、V2__Add_Vector_Knowledge.sql'],
    ['3', '模型权重', 'weather-model-api/results/model_sample.pth（ResNet-18）'],
    ['4', '部署与配置文档', 'README.md、application.yml 示例、环境变量说明'],
    ['5', '需求与设计文档', '参赛《软件应用与开发类作品设计和开发文档》'],
    ['6', '实训总结报告', '本文件'],
], )
page_break()

h1('17 总结与展望')
h2('17.1 孙瑞铭')
para('本文所做的主要工作有如下几方面：')
bullet('（1）参与系统需求分析与用例建模，梳理了天气查询、AI 气象助手、天气图像识别等核心业务的流程与功能边界；')
bullet('（2）参与后端部分模块（用户认证、天气服务、AI 智能体接口）的设计与编码，理解了 Spring Boot 分层架构与 MyBatis 持久化；')
bullet('（3）参与前端页面与可视化组件的实现，运用 Vue 3、Pinia 与 ECharts 完成天气卡片、图表与地图展示；')
bullet('（4）参与 Python 天气图像识别模型服务（Flask + ResNet-18）的对接与测试，验证了四分类推理链路；')
bullet('（5）参与系统分层测试与本文档的编写，进一步建立了软件工程全生命周期的认识。')
para('通过对 SkyGazer 系统的需求、功能与数据分析过程，我不仅深入理解了智能天气系统的业务内涵，也更加体会到软件工程方法论的价值——从需求牵引到设计落地、再到测试闭环。在技术上，我系统学习了 Spring Boot、Vue 3、Redis 缓存、Spring AI 集成大模型以及深度学习模型部署等知识，提升了工程化与协作能力。')
para('虽然本文完整研究、分析与实现了 SkyGazer 系统，但系统仍存在可改进之处：其一，天气图像识别在夜间、雾天等复杂光照条件下准确率有限；其二，RAG 检索依赖外部 embedding 端点，未开通时仅能降级；其三，当前以 Web 端为主，尚缺移动端 App；其四，多模型对比与国际化尚未开展。')
para('未来工作中，我计划在以下方向持续完善：引入 Milvus 等专业向量数据库以提升 RAG 检索质量；接入更多气象数据源（如中国气象局、NOAA）增强数据韧性；开发移动端（React Native）覆盖更多场景；完善监控告警（Prometheus + Grafana）与多语言支持，并探索 AI Agent 自主决策能力，使系统从"智能查询"进一步走向"主动决策服务"。')
page_break()

h1('附录')
para('附录 A  天气图像分类头（节选自 serve_api.py，结构与训练时一致）：')
code_block(
'class ClassificationHead(nn.Module):\n'
'    def __init__(self, in_features, num_classes, dropout=0.3):\n'
'        super().__init__()\n'
'        self.layers = nn.Sequential(\n'
'            nn.Dropout(dropout),\n'
'            nn.Linear(in_features, num_classes),\n'
'        )\n'
'\n'
'def _build_model():\n'
'    model = models.resnet18(weights=None)\n'
'    in_features = model.fc.in_features\n'
'    model.fc = ClassificationHead(in_features, 4, dropout=0.3)\n'
'    return model\n'
'# 类别顺序： cloud_type 无关，推理类别固定为 [cloudy, rainy, snowy, suny]'
)
para('附录 B  Spring AI 配置要点（已脱敏，密钥以环境变量注入，不写死于仓库）：')
code_block(
'spring:\n'
'  ai:\n'
'    openai:\n'
'      api-key: ${ALIYUN_AI_API_KEY}        # 以环境变量注入，仓库不留存明文\n'
'      base-url: ${ALIYUN_AI_OPENAI_BASE}  # 阿里云百炼兼容端点(不含/v1)\n'
'      chat:\n'
'        options:\n'
'          model: ${ALIYUN_AI_MODEL:qwen-plus}\n'
'          temperature: 0.8\n'
'      embedding:\n'
'        options:\n'
'          model: text-embedding-v2'
)
page_break()

h1('参考文献')
refs = [
    '王丹丹. 洛阳市孟津区智慧农业气象服务的现状与发展策略[J]. 棉花科学, 2025, 47(12): 53-55.',
    '陈圣劼, 王啸华, 许小龙, 等. 知识-场景耦合的江苏决策气象服务移动端应用系统研发与实践[J]. 气象科技, 2025, 53(06): 804-815.',
    '李荣, 崔智慧, 冯慧敏, 等. 暴雨天气市级决策气象服务标准化体系构建与实践[A]. 中国气象学会. 第一届城市气象服务学术交流会摘要集[C]. 郑州市气象局, 2025: 7-8.',
    '许兴华, 聂羽彗. 气候变化下人工影响天气与智慧农业决策系统的耦合机制及策略[J]. 农村科学实验, 2025, (23): 90-92.',
    '邵颖斌, 郭弘, 张立生, 等. 面向数智时代的决策气象服务研究——基于知识与数据协同驱动的范式转变与实践路径[J]. 中国防汛抗旱, 2026, 36(03): 56-61+67.',
    '王春丽. 大数据驱动的气象防雷减灾决策支持系统经济效益分析[J]. 乡镇企业导报, 2025, (20): 66-68.',
    '徐墨, 刘星辰, 齐鹭莹, 等. 基于"天擎"多时间尺度气象因子对吉林省西瓜产量进行决策树分析预测模型的研究[J]. 农业灾害研究, 2025, 15(10): 68-70.',
    '徐首利. 基于决策森林的农业气象灾害预警方法[J]. 智慧农业导刊, 2025, 5(19): 52-55+60.',
    '周苏, 王文. 软件工程学教程[M]. 北京: 科学出版社, 2003.',
    '齐治昌, 谭庆平, 宁洪. 软件工程(第二版)[M]. 北京: 高等教育出版社, 2004.',
    '王珊, 陈红. 数据库系统原理教程[M]. 北京: 清华大学出版社, 1998.',
]
for i, r in enumerate(refs, 1):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    p.paragraph_format.line_spacing = Pt(16)
    run = p.add_run('[%d] %s' % (i, r))
    set_run_font(run, size=10.5)

doc.save(OUT)
print('SAVED:', OUT)
print('paragraphs:', len(doc.paragraphs), 'tables:', len(doc.tables))
