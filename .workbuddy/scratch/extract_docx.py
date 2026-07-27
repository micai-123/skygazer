import docx
import sys

def extract(path):
    print("="*80)
    print("FILE:", path)
    print("="*80)
    doc = docx.Document(path)
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if txt:
            style = p.style.name if p.style else ""
            print(f"[{i}|{style}] {txt}")
    for ti, t in enumerate(doc.tables):
        print(f"\n--- TABLE {ti} ---")
        for r in t.rows:
            cells = [c.text.strip() for c in r.cells]
            print(" | ".join(cells))
    print("\n")

if __name__ == "__main__":
    for f in sys.argv[1:]:
        extract(f)
